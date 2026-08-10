"""
文件解析服务
支持 PDF、Word、TXT 等格式的合同文件解析
"""
import os
import re
import logging
from typing import Dict, Any, Optional, List
from io import BytesIO
import PyPDF2
from docx import Document
import pytesseract
from PIL import Image
import pdf2image
from app.core.config import settings

logger = logging.getLogger(__name__)


class ParserService:
    """文件解析服务"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'doc', 'docx', 'txt']
        
    async def parse_contract(self, file_path: str) -> Dict[str, Any]:
        """
        解析合同文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果字典
        """
        try:
            # 检查文件格式
            file_ext = os.path.splitext(file_path)[1].lower().lstrip('.')
            if file_ext not in self.supported_formats:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            # 根据文件类型调用相应的解析方法
            if file_ext == 'pdf':
                result = await self._parse_pdf(file_path)
            elif file_ext in ['doc', 'docx']:
                result = await self._parse_word(file_path)
            elif file_ext == 'txt':
                result = await self._parse_text(file_path)
            else:
                raise ValueError(f"未知的文件格式: {file_ext}")
            
            # 后处理：提取结构化信息
            structured_data = self._extract_structured_info(result['text'])
            
            return {
                "success": True,
                "text": result['text'],
                "structured_data": structured_data,
                "page_count": result.get('page_count', 1),
                "word_count": len(result['text'].split()),
                "file_format": file_ext,
                "metadata": result.get('metadata', {})
            }
            
        except Exception as e:
            logger.error(f"解析合同文件失败: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "structured_data": {},
                "page_count": 0,
                "word_count": 0,
                "file_format": file_ext if 'file_ext' in locals() else 'unknown'
            }
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析 PDF 文件"""
        text = ""
        metadata = {}
        page_count = 0
        
        try:
            # 使用 PyPDF2 解析文本
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                # 提取元数据
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                    }
                
                # 提取文本
                for page_num in range(page_count):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += f"=== 第 {page_num + 1} 页 ===\n{page_text}\n\n"
            
            # 如果文本提取失败或为空，尝试 OCR
            if not text.strip():
                logger.info("PDF 文本提取失败，尝试 OCR 识别")
                text = await self._ocr_pdf(file_path)
        
        except Exception as e:
            logger.error(f"PDF 解析失败: {e}")
            # 尝试 OCR 作为后备方案
            try:
                text = await self._ocr_pdf(file_path)
            except Exception as ocr_error:
                logger.error(f"PDF OCR 也失败: {ocr_error}")
                raise
        
        return {
            "text": text,
            "page_count": page_count,
            "metadata": metadata
        }
    
    async def _ocr_pdf(self, file_path: str) -> str:
        """使用 OCR 识别 PDF 中的文本"""
        try:
            # 尝试使用 pdf2image 将 PDF 转换为图像
            try:
                images = pdf2image.convert_from_path(file_path)
                
                text = ""
                for i, image in enumerate(images):
                    # 使用 Tesseract OCR 识别文本
                    page_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
                    text += f"=== 第 {i + 1} 页 (OCR) ===\n{page_text}\n\n"
                
                return text
            except Exception as pdf2image_error:
                logger.warning(f"pdf2image 转换失败: {pdf2image_error}")
                # 回退方案：返回简单的占位文本
                return f"[OCR 功能暂时不可用，请确保已安装 poppler]\n文件路径: {file_path}\n\n您可以手动上传文本格式的合同文件。"
            
        except Exception as e:
            logger.error(f"PDF OCR 失败: {e}")
            # 返回友好的错误信息而不是抛出异常
            return f"[OCR 处理失败: {str(e)}]\n\n建议：请确保已安装必要的依赖或上传文本格式的合同文件。"
    
    async def _parse_word(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文件"""
        try:
            doc = Document(file_path)
            
            # 提取文本
            text = ""
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += row_text + "\n"
            
            # 提取元数据
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
            }
            
            return {
                "text": text,
                "page_count": 1,  # Word 文档没有明确的页数概念
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Word 解析失败: {e}")
            raise
    
    async def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text,
                "page_count": 1,
                "metadata": {}
            }
            
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    text = file.read()
                
                return {
                    "text": text,
                    "page_count": 1,
                    "metadata": {}
                }
            except Exception as e:
                logger.error(f"文本文件解析失败: {e}")
                raise
    
    def _extract_structured_info(self, text: str) -> Dict[str, Any]:
        """从文本中提取结构化信息"""
        structured_info = {
            "parties": self._extract_parties(text),
            "dates": self._extract_dates(text),
            "amounts": self._extract_amounts(text),
            "clauses": self._identify_clauses(text),
            "key_terms": self._extract_key_terms(text),
        }
        
        return structured_info
    
    def _extract_parties(self, text: str) -> List[Dict[str, str]]:
        """提取合同当事人信息"""
        parties = []
        
        # 常见当事人模式
        patterns = [
            r'(甲方|发包方|买方|出租方|委托方)[：:]\s*([^\n]+)',
            r'(乙方|承包方|卖方|承租方|受托方)[：:]\s*([^\n]+)',
            r'(丙方)[：:]\s*([^\n]+)',
            r'([^\n]+)\s*\(以下简称["「]?甲方["」]?\)',
            r'([^\n]+)\s*\(以下简称["「]?乙方["」]?\)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                if len(match.groups()) >= 2:
                    role = match.group(1)
                    name = match.group(2).strip()
                    parties.append({
                        "role": role,
                        "name": name,
                        "type": self._classify_party_type(role)
                    })
        
        return parties
    
    def _classify_party_type(self, role: str) -> str:
        """分类当事人类型"""
        buyer_roles = ['甲方', '买方', '采购方', '委托方']
        seller_roles = ['乙方', '卖方', '供应方', '受托方']
        
        if role in buyer_roles:
            return "buyer"
        elif role in seller_roles:
            return "seller"
        else:
            return "other"
    
    def _extract_dates(self, text: str) -> List[Dict[str, str]]:
        """提取日期信息"""
        dates = []
        
        # 日期模式
        patterns = [
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
            r'(\d{4})-(\d{1,2})-(\d{1,2})',
            r'(\d{4})/(\d{1,2})/(\d{1,2})',
            r'(\d{1,2})月(\d{1,2})日',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                date_str = match.group(0)
                dates.append({
                    "date": date_str,
                    "context": self._get_context(text, match.start(), match.end())
                })
        
        return dates
    
    def _extract_amounts(self, text: str) -> List[Dict[str, Any]]:
        """提取金额信息"""
        amounts = []
        
        # 金额模式
        patterns = [
            r'(人民币|RMB|CNY|¥|￥)\s*([0-9,]+(?:\.\d{1,2})?)',
            r'([0-9,]+(?:\.\d{1,2})?)\s*(万元|元|人民币|RMB)',
            r'总价[：:]\s*([0-9,]+(?:\.\d{1,2})?)',
            r'金额[：:]\s*([0-9,]+(?:\.\d{1,2})?)',
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                amount_str = match.group(0)
                amounts.append({
                    "amount": amount_str,
                    "context": self._get_context(text, match.start(), match.end())
                })
        
        return amounts
    
    def _identify_clauses(self, text: str) -> List[Dict[str, Any]]:
        """识别合同条款"""
        clauses = []
        
        # 常见条款标题
        clause_titles = [
            ("第一条", "合同双方"),
            ("第二条", "合同标的"),
            ("第三条", "价款及支付"),
            ("第四条", "交付及验收"),
            ("第五条", "质量保证"),
            ("第六条", "违约责任"),
            ("第七条", "争议解决"),
            ("第八条", "保密条款"),
            ("第九条", "知识产权"),
            ("第十条", "其他约定"),
            
            ("第1条", "合同双方"),
            ("第2条", "合同标的"),
            ("第3条", "价款及支付"),
            ("第4条", "交付及验收"),
            ("第5条", "质量保证"),
            ("第6条", "违约责任"),
            ("第7条", "争议解决"),
            ("第8条", "保密条款"),
            ("第9条", "知识产权"),
            ("第10条", "其他约定"),
            
            ("一、", "合同双方"),
            ("二、", "合同标的"),
            ("三、", "价款及支付"),
            ("四、", "交付及验收"),
            ("五、", "质量保证"),
            ("六、", "违约责任"),
            ("七、", "争议解决"),
            ("八、", "保密条款"),
            ("九、", "知识产权"),
            ("十、", "其他约定"),
            
            ("1.", "合同双方"),
            ("2.", "合同标的"),
            ("3.", "价款及支付"),
            ("4.", "交付及验收"),
            ("5.", "质量保证"),
            ("6.", "违约责任"),
            ("7.", "争议解决"),
            ("8.", "保密条款"),
            ("9.", "知识产权"),
            ("10.", "其他约定"),
        ]
        
        for pattern, default_title in clause_titles:
            matches = list(re.finditer(rf'^{pattern}\s*([^\n]*)', text, re.MULTILINE))
            for match in matches:
                title = match.group(1).strip() or default_title
                start_pos = match.start()
                
                # 查找条款结束位置（下一个条款开始或文档结束）
                end_pos = len(text)
                next_match = re.search(rf'^{pattern[0] if pattern[0].isdigit() else pattern}', 
                                      text[start_pos+1:], re.MULTILINE)
                if next_match:
                    end_pos = start_pos + next_match.start()
                
                clause_text = text[start_pos:end_pos].strip()
                clauses.append({
                    "title": title,
                    "content": clause_text,
                    "start_pos": start_pos,
                    "end_pos": end_pos
                })
        
        return clauses
    
    def _extract_key_terms(self, text: str) -> List[str]:
        """提取关键词"""
        key_terms = []
        
        # 法律和合同相关关键词
        legal_terms = [
            "违约责任", "争议解决", "仲裁", "诉讼", "管辖权",
            "保密义务", "知识产权", "不可抗力", "解除合同",
            "违约金", "赔偿", "担保", "保证", "质押", "抵押",
            "生效", "终止", "续约", "变更", "转让",
        ]
        
        for term in legal_terms:
            if term in text:
                key_terms.append(term)
        
        return list(set(key_terms))  # 去重
    
    def _get_context(self, text: str, start: int, end: int, context_len: int = 100) -> str:
        """获取上下文文本"""
        context_start = max(0, start - context_len)
        context_end = min(len(text), end + context_len)
        return text[context_start:context_end]


# 创建全局解析服务实例
parser_service = ParserService()